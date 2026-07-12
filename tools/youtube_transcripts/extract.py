#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
EXIT:MARU — 유튜브 영상 자막(글자) 추출 → 영상별 저장 + 워드 파일 1개

이 스크립트가 하는 일 (쉽게):
  1) youtube_urls.txt 에 적힌 유튜브 주소를 한 줄씩 읽는다 (재생목록이면 안의 영상들로 펼친다).
  2) 영상마다 '자막(사람이 쓴 글자, 없으면 자동 생성 자막)'을 통째로 가져온다.
  3) 영상 1개당 텍스트 파일 1개로 저장한다  →  output/텍스트/0001_영상제목.txt
  4) 마지막에 전부 묶은 워드 파일 1개를 만든다  →  output/유튜브_자막_모음.docx

★ 왜 이걸 '내 컴퓨터'에서 돌려야 하나?
  유튜브는 클라우드 서버(데이터센터) IP에서 오는 자동 자막 요청을 봇으로 보고 막는다.
  집/회사 인터넷(일반 IP)에서는 대부분 정상 동작한다. 그래서 이 파일을 받아
  네 컴퓨터에서 실행하는 것이다.  (설치·실행법은 같은 폴더 README.md 참고)

필요 패키지:  pip install youtube-transcript-api python-docx yt-dlp
  - youtube-transcript-api : 자막 가져오기 (핵심)
  - python-docx            : 워드(.docx) 만들기
  - yt-dlp                 : 재생목록 펼치기 + 자막 백업 경로 (없어도 개별영상은 동작)
"""

import os
import re
import sys
import time
import json
import html
import urllib.request

# ───────────────────────── 설정 (원하면 여기만 바꾸면 됨) ─────────────────────────
HERE = os.path.dirname(os.path.abspath(__file__))
URL_FILE = os.path.join(HERE, "youtube_urls.txt")     # 입력: 유튜브 주소 목록
OUT_DIR = os.path.join(HERE, "output")                # 출력 폴더
TXT_DIR = os.path.join(OUT_DIR, "텍스트")              # 영상별 텍스트 저장 위치
DOCX_PATH = os.path.join(OUT_DIR, "유튜브_자막_모음.docx")
LANG_PREF = ["ko", "ko-KR", "en", "en-US"]            # 우선순위: 한국어 먼저, 없으면 영어
SLEEP_SEC = 1.0                                        # 영상 사이 쉬는 시간(초) — 차단 회피
RETRY = 2                                              # 실패 시 재시도 횟수
TIMESTAMPS = False                                     # True면 각 줄 앞에 [분:초] 시간 표시
# ──────────────────────────────────────────────────────────────────────────────


def log(msg):
    print(msg, flush=True)


def read_input_lines():
    if not os.path.exists(URL_FILE):
        log(f"[오류] 입력 파일이 없습니다: {URL_FILE}")
        sys.exit(1)
    out = []
    for line in open(URL_FILE, encoding="utf-8"):
        s = line.strip()
        if not s or s.startswith("#"):
            continue
        out.append(s)
    return out


def extract_video_id(url):
    """watch?v= / youtu.be/ / shorts/ 형태에서 11자리 영상 ID를 뽑는다."""
    m = re.search(r"(?:watch\?v=|youtu\.be/|shorts/|/embed/)([A-Za-z0-9_-]{11})", url)
    if m:
        return m.group(1)
    # 그냥 ID만 적힌 줄도 허용
    if re.fullmatch(r"[A-Za-z0-9_-]{11}", url):
        return url
    return None


def extract_playlist_id(url):
    m = re.search(r"list=(PL[A-Za-z0-9_-]+)", url)
    return m.group(1) if m else None


def expand_playlist(playlist_id):
    """yt-dlp로 재생목록 안의 영상 ID들을 펼친다. yt-dlp 없으면 빈 목록 반환."""
    try:
        import yt_dlp
    except ImportError:
        log(f"  [건너뜀] yt-dlp가 없어 재생목록을 펼칠 수 없음: {playlist_id}")
        return []
    ids = []
    opts = {"quiet": True, "no_warnings": True, "extract_flat": True, "skip_download": True}
    try:
        with yt_dlp.YoutubeDL(opts) as ydl:
            info = ydl.extract_info(f"https://www.youtube.com/playlist?list={playlist_id}", download=False)
            for e in info.get("entries", []) or []:
                vid = e.get("id")
                if vid and re.fullmatch(r"[A-Za-z0-9_-]{11}", vid):
                    ids.append(vid)
        log(f"  재생목록 {playlist_id}: {len(ids)}개 영상")
    except Exception as e:
        log(f"  [실패] 재생목록 {playlist_id}: {type(e).__name__} {str(e)[:120]}")
    return ids


def get_title(video_id):
    """oEmbed로 제목/채널을 가볍게 가져온다 (API 키 불필요). 실패하면 ID로 대체."""
    url = f"https://www.youtube.com/oembed?url=https://www.youtube.com/watch?v={video_id}&format=json"
    try:
        req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
        data = json.loads(urllib.request.urlopen(req, timeout=15).read().decode("utf-8"))
        return data.get("title") or video_id, data.get("author_name") or ""
    except Exception:
        return video_id, ""


def _fmt_ts(sec):
    sec = int(sec)
    return f"[{sec // 60:02d}:{sec % 60:02d}] "


def fetch_transcript(video_id):
    """
    자막을 가져온다. 반환: (텍스트, 사용언어, 방법) 또는 (None, None, 실패사유).
    youtube-transcript-api 신/구 버전 모두 대응. 실패 시 yt-dlp 백업 경로 시도.
    """
    # ── 1순위: youtube-transcript-api ──
    try:
        from youtube_transcript_api import YouTubeTranscriptApi
        snippets = None
        used_lang = None

        # 신 인터페이스 (v1.x): 인스턴스.fetch()
        try:
            api = YouTubeTranscriptApi()
            fetched = api.fetch(video_id, languages=LANG_PREF)
            snippets = [(s.text, s.start) for s in fetched]
            used_lang = getattr(fetched, "language_code", "?")
        except TypeError:
            # 구 인터페이스 (v0.6.x): 클래스 정적 메서드
            data = YouTubeTranscriptApi.get_transcript(video_id, languages=LANG_PREF)
            snippets = [(d["text"], d.get("start", 0)) for d in data]
            used_lang = LANG_PREF[0]
        except Exception:
            # 우선언어가 없을 때: 사용 가능한 자막 아무거나 (자동생성 포함, 한국어로 번역 시도)
            try:
                api = YouTubeTranscriptApi()
                tlist = api.list(video_id)
                tr = None
                try:
                    tr = tlist.find_transcript(LANG_PREF)
                except Exception:
                    for t in tlist:
                        tr = t
                        break
                if tr is not None:
                    try:
                        if tr.language_code not in ("ko", "ko-KR") and tr.is_translatable:
                            tr = tr.translate("ko")
                    except Exception:
                        pass
                    f2 = tr.fetch()
                    snippets = [(s.text, s.start) for s in f2]
                    used_lang = getattr(tr, "language_code", "?")
            except Exception:
                snippets = None

        if snippets:
            if TIMESTAMPS:
                text = "\n".join(_fmt_ts(st) + html.unescape(t).replace("\n", " ") for t, st in snippets)
            else:
                text = "\n".join(html.unescape(t).replace("\n", " ") for t, _ in snippets)
            return text.strip(), used_lang, "youtube-transcript-api"
    except ImportError:
        pass
    except Exception:
        pass

    # ── 2순위: yt-dlp 자막 백업 ──
    try:
        text, lang = fetch_transcript_ytdlp(video_id)
        if text:
            return text, lang, "yt-dlp"
    except Exception:
        pass

    return None, None, "자막 없음/차단"


def fetch_transcript_ytdlp(video_id):
    """yt-dlp로 자막 트랙 URL을 얻어 timedtext(json3)를 직접 받아 텍스트로 변환."""
    import yt_dlp
    opts = {"skip_download": True, "quiet": True, "no_warnings": True,
            "writesubtitles": True, "writeautomaticsub": True}
    with yt_dlp.YoutubeDL(opts) as ydl:
        info = ydl.extract_info(f"https://www.youtube.com/watch?v={video_id}", download=False)
    subs = info.get("subtitles") or {}
    auto = info.get("automatic_captions") or {}
    track = None
    lang = None
    for L in LANG_PREF:
        if L in subs:
            track, lang = subs[L], L; break
    if track is None:
        for L in LANG_PREF:
            if L in auto:
                track, lang = auto[L], L; break
    if track is None:
        # 아무 자막이나
        for src in (subs, auto):
            for L, tr in src.items():
                track, lang = tr, L; break
            if track:
                break
    if not track:
        return None, None
    fmt = next((f for f in track if f.get("ext") == "json3"), track[0])
    req = urllib.request.Request(fmt["url"], headers={"User-Agent": "Mozilla/5.0"})
    raw = urllib.request.urlopen(req, timeout=25).read().decode("utf-8", "ignore")
    lines = []
    try:
        j = json.loads(raw)
        for ev in j.get("events", []):
            seg = "".join(s.get("utf8", "") for s in ev.get("segs", []) or [])
            seg = seg.replace("\n", " ").strip()
            if seg:
                lines.append(seg)
    except json.JSONDecodeError:
        # vtt/srv 형태면 태그 제거
        raw = re.sub(r"<[^>]+>", "", raw)
        for ln in raw.splitlines():
            ln = ln.strip()
            if ln and "-->" not in ln and not ln.isdigit() and ln != "WEBVTT":
                lines.append(ln)
    return ("\n".join(lines).strip() or None), lang


def safe_filename(name, maxlen=60):
    name = re.sub(r'[\\/:*?"<>|\n\r\t]', "_", name).strip()
    return (name[:maxlen]).strip() or "untitled"


def main():
    os.makedirs(TXT_DIR, exist_ok=True)

    # 1) 입력 → 영상 ID 목록(중복 제거, 순서 유지) + 재생목록 펼치기
    log("입력 목록 읽는 중...")
    video_ids = []
    seen = set()

    def add(vid):
        if vid and vid not in seen:
            seen.add(vid); video_ids.append(vid)

    for line in read_input_lines():
        pl = extract_playlist_id(line)
        if pl and "watch?v=" not in line:
            for vid in expand_playlist(pl):
                add(vid)
            continue
        add(extract_video_id(line))

    log(f"총 {len(video_ids)}개 영상 처리 시작\n")

    # 2) 워드 문서 준비
    from docx import Document
    from docx.shared import Pt, RGBColor
    doc = Document()
    doc.add_heading("유튜브 자막 모음 — EXIT:MARU", level=0)
    intro = doc.add_paragraph()
    intro.add_run(f"총 {len(video_ids)}개 영상 · youtube-transcript-api 추출").italic = True
    doc.add_paragraph("")

    results = []  # (idx, vid, title, author, ok, method, lang, nchars)

    for i, vid in enumerate(video_ids, 1):
        title, author = get_title(vid)
        text, lang, method = (None, None, None)
        for attempt in range(RETRY + 1):
            text, lang, method = fetch_transcript(vid)
            if text:
                break
            time.sleep(1.5 * (attempt + 1))  # 실패 시 점점 더 쉼

        ok = bool(text)
        nchars = len(text) if ok else 0
        results.append((i, vid, title, author, ok, method, lang, nchars))

        # 영상별 텍스트 파일 저장
        fname = f"{i:04d}_{safe_filename(title)}.txt"
        with open(os.path.join(TXT_DIR, fname), "w", encoding="utf-8") as f:
            f.write(f"제목: {title}\n채널: {author}\n주소: https://www.youtube.com/watch?v={vid}\n")
            f.write(f"언어: {lang or '-'} · 방법: {method}\n{'='*50}\n\n")
            f.write(text if ok else "[자막을 가져오지 못함: 자막 비공개이거나 요청 차단]")

        # 워드에 한 섹션 추가
        doc.add_heading(f"{i}. {title}", level=1)
        meta = doc.add_paragraph()
        meta.add_run(f"https://www.youtube.com/watch?v={vid}").font.size = Pt(9)
        meta.add_run(f"   ·  {author}").font.size = Pt(9)
        if ok:
            for para in text.split("\n"):
                if para.strip():
                    doc.add_paragraph(para.strip())
        else:
            warn = doc.add_paragraph()
            r = warn.add_run("⚠ 자막을 가져오지 못했습니다 (자막 비공개 또는 요청 차단).")
            r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        doc.add_page_break()

        mark = "OK " if ok else "실패"
        log(f"[{i:>3}/{len(video_ids)}] {mark} {vid} · {lang or '-':5} · {nchars:>6}자 · {title[:40]}")
        time.sleep(SLEEP_SEC)

    doc.save(DOCX_PATH)

    # 3) 요약 리포트
    ok_n = sum(1 for r in results if r[4])
    log("\n" + "=" * 50)
    log(f"완료: 성공 {ok_n} / 전체 {len(video_ids)}")
    log(f"워드 파일 : {DOCX_PATH}")
    log(f"영상별 txt: {TXT_DIR}")
    # 실패 목록 CSV
    with open(os.path.join(OUT_DIR, "결과요약.csv"), "w", encoding="utf-8-sig") as f:
        f.write("번호,영상ID,제목,채널,성공,방법,언어,글자수\n")
        for (i, vid, title, author, ok, method, lang, n) in results:
            t = title.replace(",", " ")
            a = author.replace(",", " ")
            f.write(f"{i},{vid},{t},{a},{'성공' if ok else '실패'},{method},{lang or ''},{n}\n")
    log(f"요약 CSV  : {os.path.join(OUT_DIR, '결과요약.csv')}")


if __name__ == "__main__":
    main()
